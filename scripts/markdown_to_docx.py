import re
import sys
from pathlib import Path

from docx import Document


def flush_list(document: Document, list_buffer, numbered: bool) -> None:
    style = "List Number" if numbered else "List Bullet"
    for item in list_buffer:
        document.add_paragraph(item, style=style)


def convert_markdown_to_docx(md_path: Path, docx_path: Path) -> None:
    if not md_path.exists():
        raise FileNotFoundError(f"Markdown file not found: {md_path}")

    lines = md_path.read_text(encoding="utf-8").splitlines()
    document = Document()

    bullet_buffer = []
    number_buffer = []

    def flush_all_lists() -> None:
        nonlocal bullet_buffer, number_buffer
        if bullet_buffer:
            flush_list(document, bullet_buffer, numbered=False)
            bullet_buffer = []
        if number_buffer:
            flush_list(document, number_buffer, numbered=True)
            number_buffer = []

    for raw_line in lines:
        line = raw_line.rstrip()

        if not line.strip():
            flush_all_lists()
            continue

        heading_match = re.match(r"^(#{1,6})\s+(.*)$", line)
        if heading_match:
            flush_all_lists()
            level = len(heading_match.group(1))
            text = heading_match.group(2).strip()
            document.add_heading(text, level=min(level, 4))
            continue

        bullet_match = re.match(r"^[-*]\s+(.*)$", line)
        if bullet_match:
            if number_buffer:
                flush_list(document, number_buffer, numbered=True)
                number_buffer = []
            bullet_buffer.append(bullet_match.group(1).strip())
            continue

        number_match = re.match(r"^\d+\.\s+(.*)$", line)
        if number_match:
            if bullet_buffer:
                flush_list(document, bullet_buffer, numbered=False)
                bullet_buffer = []
            number_buffer.append(number_match.group(1).strip())
            continue

        flush_all_lists()
        document.add_paragraph(line)

    flush_all_lists()

    docx_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(docx_path)


if __name__ == "__main__":
    if len(sys.argv) != 3:
        print("Usage: python scripts/markdown_to_docx.py <input.md> <output.docx>")
        sys.exit(1)

    input_md = Path(sys.argv[1])
    output_docx = Path(sys.argv[2])
    convert_markdown_to_docx(input_md, output_docx)
    print(f"Created {output_docx}")
