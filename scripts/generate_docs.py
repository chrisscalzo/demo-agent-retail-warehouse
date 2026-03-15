"""
generate_docs.py
Contoso Retail Warehouse Demo — Document Generator

Walks content/ directory, reads each .json file, and produces a styled
Word document (.docx) in the matching output/ folder structure.

Usage:
    python generate_docs.py

Dependencies:
    pip install python-docx
"""

import json
import os
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor, Inches

# ── Configuration ────────────────────────────────────────────────────────────

CONTENT_DIR = Path("content")
OUTPUT_DIR = Path("output")

# Contoso brand colours
BRAND_DARK_BLUE = RGBColor(0x00, 0x3A, 0x70)   # header background
BRAND_ORANGE = RGBColor(0xF2, 0x6B, 0x00)       # accent
BRAND_LIGHT_GREY = RGBColor(0xF5, 0xF5, 0xF5)  # table header shading
WHITE = RGBColor(0xFF, 0xFF, 0xFF)

# Content-dir → output-dir sub-path mapping
PATH_MAP = {
    Path("content/general"): Path("output/General"),
    Path("content/states/ca"): Path("output/States/CA"),
    Path("content/states/nv"): Path("output/States/NV"),
    Path("content/states/az"): Path("output/States/AZ"),
}


# ── Helpers ──────────────────────────────────────────────────────────────────

def set_cell_bg(cell, hex_color: str) -> None:
    """Set a table cell's background shading colour."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def hex_from_rgb(rgb: RGBColor) -> str:
    return f"{rgb[0]:02X}{rgb[1]:02X}{rgb[2]:02X}"


def bold_run(para, text: str) -> None:
    run = para.add_run(text)
    run.bold = True


def add_horizontal_rule(doc: Document) -> None:
    """Insert a thin horizontal rule paragraph."""
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(2)
    para.paragraph_format.space_after = Pt(2)
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), f"{hex_from_rgb(BRAND_ORANGE)}")
    pBdr.append(bottom)
    pPr.append(pBdr)


# ── Document building ─────────────────────────────────────────────────────────

def build_document(data: dict) -> Document:
    doc = Document()

    # ── Page margins ────────────────────────────────────────────────────────
    for section in doc.sections:
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)

    # ── Header block ─────────────────────────────────────────────────────────
    company_para = doc.add_paragraph()
    company_para.paragraph_format.space_before = Pt(0)
    company_para.paragraph_format.space_after = Pt(4)
    run = company_para.add_run("CONTOSO RETAIL")
    run.bold = True
    run.font.size = Pt(9)
    run.font.color.rgb = BRAND_ORANGE
    run.font.all_caps = True

    title_para = doc.add_paragraph()
    title_para.paragraph_format.space_before = Pt(0)
    title_para.paragraph_format.space_after = Pt(2)
    run = title_para.add_run(data["title"])
    run.bold = True
    run.font.size = Pt(20)
    run.font.color.rgb = BRAND_DARK_BLUE

    meta_para = doc.add_paragraph()
    meta_para.paragraph_format.space_before = Pt(0)
    meta_para.paragraph_format.space_after = Pt(6)
    run = meta_para.add_run(
        f"Version {data['version']}  |  Category: {data['category']}  |  "
        f"Audience: {', '.join(data['audience'])}"
    )
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x60, 0x60, 0x60)
    run.italic = True

    add_horizontal_rule(doc)

    # ── Sections ─────────────────────────────────────────────────────────────
    for section_data in data.get("sections", []):
        heading_text = section_data.get("heading", "")
        level = section_data.get("level", 1)

        # Heading
        if level == 1:
            h = doc.add_heading(heading_text, level=1)
            h.paragraph_format.space_before = Pt(12)
            h.paragraph_format.space_after = Pt(4)
            for run in h.runs:
                run.font.color.rgb = BRAND_DARK_BLUE
                run.font.size = Pt(13)
        else:
            h = doc.add_heading(heading_text, level=2)
            h.paragraph_format.space_before = Pt(8)
            h.paragraph_format.space_after = Pt(3)
            for run in h.runs:
                run.font.color.rgb = BRAND_ORANGE
                run.font.size = Pt(11)

        # Content blocks
        for block in section_data.get("content", []):
            btype = block.get("type")

            if btype == "paragraph":
                p = doc.add_paragraph(block["text"])
                p.style = doc.styles["Normal"]
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after = Pt(6)
                for run in p.runs:
                    run.font.size = Pt(10)

            elif btype == "bullet_list":
                for item in block.get("items", []):
                    p = doc.add_paragraph(item, style="List Bullet")
                    p.paragraph_format.space_before = Pt(0)
                    p.paragraph_format.space_after = Pt(2)
                    for run in p.runs:
                        run.font.size = Pt(10)

            elif btype == "numbered_list":
                for item in block.get("items", []):
                    p = doc.add_paragraph(item, style="List Number")
                    p.paragraph_format.space_before = Pt(0)
                    p.paragraph_format.space_after = Pt(2)
                    for run in p.runs:
                        run.font.size = Pt(10)

            elif btype == "table":
                headers = block.get("headers", [])
                rows = block.get("rows", [])
                col_count = len(headers)
                if col_count == 0:
                    continue

                tbl = doc.add_table(rows=1 + len(rows), cols=col_count)
                tbl.style = "Table Grid"

                # Header row
                hdr_row = tbl.rows[0]
                for i, hdr_text in enumerate(headers):
                    cell = hdr_row.cells[i]
                    cell.text = hdr_text
                    set_cell_bg(cell, hex_from_rgb(BRAND_DARK_BLUE))
                    for para in cell.paragraphs:
                        for run in para.runs:
                            run.bold = True
                            run.font.color.rgb = WHITE
                            run.font.size = Pt(9)

                # Data rows
                for r_idx, row_data in enumerate(rows):
                    data_row = tbl.rows[r_idx + 1]
                    for c_idx, cell_text in enumerate(row_data):
                        cell = data_row.cells[c_idx]
                        cell.text = cell_text
                        # Alternate row shading
                        if r_idx % 2 == 1:
                            set_cell_bg(cell, hex_from_rgb(BRAND_LIGHT_GREY))
                        for para in cell.paragraphs:
                            for run in para.runs:
                                run.font.size = Pt(9)

                # Space after table
                doc.add_paragraph().paragraph_format.space_after = Pt(4)

    # ── Footer ───────────────────────────────────────────────────────────────
    add_horizontal_rule(doc)
    footer_para = doc.add_paragraph(
        "Contoso Retail — Internal Use Only  |  "
        "Review annually or when regulations change  |  "
        "Questions? Contact your Warehouse Manager or Safety Coordinator."
    )
    footer_para.paragraph_format.space_before = Pt(4)
    for run in footer_para.runs:
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)
        run.italic = True

    return doc


# ── Output path resolution ────────────────────────────────────────────────────

def resolve_output_path(json_path: Path) -> Path:
    """Map a content/*.json path to an output/*.docx path."""
    for src, dst in PATH_MAP.items():
        try:
            relative = json_path.relative_to(src)
            out_path = dst / relative.with_suffix(".docx")
            return out_path
        except ValueError:
            continue
    # Fallback: mirror under output/
    rel = json_path.relative_to(CONTENT_DIR)
    return OUTPUT_DIR / rel.with_suffix(".docx")


# ── Main ─────────────────────────────────────────────────────────────────────

def main() -> None:
    json_files = sorted(CONTENT_DIR.rglob("*.json"))

    if not json_files:
        print("No JSON files found in content/. Exiting.")
        return

    generated = []
    errors = []

    for json_path in json_files:
        try:
            with open(json_path, encoding="utf-8") as f:
                data = json.load(f)

            out_path = resolve_output_path(json_path)
            out_path.parent.mkdir(parents=True, exist_ok=True)

            doc = build_document(data)
            doc.save(out_path)
            generated.append(out_path)
            print(f"  [OK] {out_path}")

        except Exception as exc:
            errors.append((json_path, exc))
            print(f"  [ERROR] {json_path}: {exc}")

    print()
    print(f"{'='*60}")
    print(f"Generated: {len(generated)} documents")
    if errors:
        print(f"Errors:    {len(errors)}")
        for path, err in errors:
            print(f"  - {path}: {err}")
    print(f"Output directory: {OUTPUT_DIR.resolve()}")
    print(f"{'='*60}")


if __name__ == "__main__":
    main()
